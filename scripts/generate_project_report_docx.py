"""Generate a detailed Word documentation of the AI Study Hub project work."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "AI-StudyHub-Complete-Project-Report.docx"


def set_heading_color(run, rgb=(0x0B, 0x17, 0x24)):
    run.font.color.rgb = RGBColor(*rgb)


def add_hr(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1FAA8A")
    pBdr.append(bottom)
    pPr.append(pBdr)


def h(doc, text, level=1):
    para = doc.add_heading(text, level=level)
    color = (0x0F, 0x6F, 0x5C) if level == 1 else (0x0B, 0x17, 0x24)
    for run in para.runs:
        set_heading_color(run, color)
    return para


def p(doc, text, bold=False):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    return para


def bullets(doc, items):
    for item in items:
        para = doc.add_paragraph(item, style="List Bullet")
        for run in para.runs:
            run.font.size = Pt(11)


def numbered(doc, items):
    for item in items:
        para = doc.add_paragraph(item, style="List Number")
        for run in para.runs:
            run.font.size = Pt(11)


def table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, name in enumerate(headers):
        hdr[i].text = name
        for paragraph in hdr[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = t.rows[r_idx + 1].cells[c_idx]
            cell.text = str(value)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()


def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("AI Study Hub")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor(0x0F, 0x6F, 0x5C)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Complete Project Report — End-to-End Delivery Documentation")
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x5B, 0x6B, 0x7C)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run(
        "Repository: github.com/Waleedcoderarch/AI-StudyPlan\n"
        "AWS Region: us-east-1  |  Account: 588681235213\n"
        "Document generated: July 2026"
    )
    r.font.size = Pt(10)
    add_hr(doc.add_paragraph())

    # 1
    h(doc, "1. Executive Summary", 1)
    p(
        doc,
        "This report documents the full journey of taking an existing AI Study Hub application "
        "(doubt solver, PDF notes generator, and quiz generator) and transforming it into a "
        "cloud-ready platform on Amazon Web Services. The work covered Infrastructure as Code, "
        "containerization, multi-environment deployment, CI/CD, cost controls, observability "
        "(Prometheus/Grafana), and a redesigned frontend UI/UX."
    )
    p(doc, "Primary outcomes:", bold=True)
    bullets(
        doc,
        [
            "Application deployed on AWS ECS Fargate behind an Application Load Balancer in us-east-1.",
            "Terraform modules and environments for dev, staging, and production (prod prepared, gated).",
            "GitHub Actions CI/CD: auto-deploy to DEV, promote to STAGING, manual confirm for PROD.",
            "Cost budgets and SNS email alerts for spend visibility.",
            "Staging observability stack: YACE + Prometheus + Grafana with scaling/reliability alarms.",
            "Frontend redesigned to a brand-first 3D “study plane” experience with intentional motion.",
            "End-to-end documentation added to the repository.",
        ],
    )

    # 2
    h(doc, "2. Starting Point — What Already Existed", 1)
    p(
        doc,
        "The product codebase already existed under Waleedcoderarch/AI-StudyPlan. It was a full-stack "
        "student assistant with three features: AI Doubt Solver, Notes Generator (PDF), and Quiz Generator. "
        "The frontend used React + Vite + Tailwind; the backend used Node.js + Express; AI responses were "
        "powered through an OpenAI-compatible client (later wired to Groq). Persistence used SQLite via Sequelize."
    )
    p(doc, "Important principle: infrastructure was added in the same repository without rewriting the core app. "
          "Folders such as infra/, docker/, and .github/ were added alongside client/ and server/.")

    # 3
    h(doc, "3. Project Setup & Accounts", 1)
    h(doc, "3.1 Local workspace", 2)
    numbered(
        doc,
        [
            "Created Projects directory under the user profile.",
            "Cloned https://github.com/Waleedcoderarch/AI-StudyPlan.git to "
            "C:\\Users\\IRUM NAUREEN\\Projects\\AI-StudyPlan.",
            "Moved the Cursor agent workspace into that project root for all subsequent work.",
        ],
    )
    h(doc, "3.2 GitHub", 2)
    bullets(
        doc,
        [
            "Target GitHub account: Waleedcoderarch (confirmed by the user).",
            "GitHub CLI authenticated as Waleedcoderarch.",
            "Remote origin remained https://github.com/Waleedcoderarch/AI-StudyPlan.git.",
            "All infra, CI/CD, monitoring, frontend, and docs commits were pushed to main on this repo.",
            "Two-account management was discussed (gh auth login / gh auth switch); work proceeded on Waleedcoderarch.",
        ],
    )
    h(doc, "3.3 AWS", 2)
    bullets(
        doc,
        [
            "AWS CLI configured for region us-east-1.",
            "Verified identity with aws sts get-caller-identity (IAM user new-project on account 588681235213).",
            "Credentials were configured locally via aws configure — never committed to git.",
            "User was advised not to paste access keys into chat.",
        ],
    )

    # 4
    h(doc, "4. Architecture Decisions", 1)
    h(doc, "4.1 ECS Fargate instead of EKS / Minikube", 2)
    p(
        doc,
        "Initial planning discussed Minikube (local) and EKS (AWS Kubernetes). The final cloud path chosen "
        "was Amazon ECS on Fargate for lower operational overhead and better cost profile for this stage."
    )
    table(
        doc,
        ["Concept", "Kubernetes", "This project"],
        [
            ["Workload unit", "Pod", "Fargate task"],
            ["Replica controller", "Deployment", "ECS Service + Task Definition"],
            ["Service discovery / ingress", "Service + Ingress", "ALB + target groups + path rules"],
            ["Cluster", "EKS / Minikube", "ECS cluster"],
            ["Node autoscaling", "Cluster Autoscaler / ASG", "Not used (no EC2 nodes)"],
            ["App autoscaling", "HPA", "ECS Application Auto Scaling (CPU)"],
        ],
    )
    h(doc, "4.2 Request routing", 2)
    p(doc, "Internet → Application Load Balancer → path-based forwarding:")
    bullets(
        doc,
        [
            "/api/* and /uploads/* → server container port 5000",
            "/* → client nginx container port 80",
            "/grafana and /grafana/* → Grafana (staging monitoring stack)",
        ],
    )
    h(doc, "4.3 Data & secrets", 2)
    bullets(
        doc,
        [
            "SQLite file stored on EFS so data survives task replacement.",
            "Groq API key stored in AWS Secrets Manager.",
            "App expects OPENAI_API_KEY with OPENAI_BASE_URL=https://api.groq.com/openai/v1 "
            "(OpenAI-compatible SDK talking to Groq).",
            "CORS_ORIGIN set to the ALB URL so the browser frontend can call the API.",
        ],
    )

    # 5
    h(doc, "5. Infrastructure as Code (Terraform)", 1)
    p(doc, "All AWS resources were expressed as Terraform modules and environment stacks under infra/.")
    h(doc, "5.1 Modules", 2)
    table(
        doc,
        ["Module", "Responsibility"],
        [
            ["vpc", "VPC, public/private subnets, Internet Gateway; NAT disabled by default for cost"],
            ["ecr", "ECR repositories for images; lifecycle keep last 5 tags"],
            ["alb", "Public ALB, security group, server/client target groups, HTTP listener rules"],
            ["efs", "Encrypted EFS + access point for SQLite persistence"],
            ["ecs", "Cluster, task defs, Fargate services, IAM roles, CloudWatch logs, CPU autoscaling"],
            ["cost_alerts", "AWS Budgets, SNS topic/email, CloudWatch estimated charges alarm"],
            ["monitoring", "Grafana ALB rule, monitoring ECS task, CloudWatch scaling alarms"],
        ],
    )
    h(doc, "5.2 Environments", 2)
    table(
        doc,
        ["Environment", "Status", "Intent"],
        [
            ["dev", "Applied & live", "Daily development; CI auto-deploy; small Fargate sizes"],
            ["staging", "Applied & live", "Pre-production testing + Grafana observability"],
            ["prod", "Code ready, not applied", "Larger capacity; deploy only after staging + human confirm"],
        ],
    )
    h(doc, "5.3 Cost-aware defaults", 2)
    bullets(
        doc,
        [
            "No NAT Gateway in default layouts (significant monthly savings).",
            "Small Fargate CPU/memory in dev/staging (e.g. 256 CPU / 512 MB class sizes where configured).",
            "Short CloudWatch log retention.",
            "ECR lifecycle policies to limit stored images.",
            "Budgets: roughly $30/mo for dev and $50/mo for staging (adjust after real Cost Explorer data).",
        ],
    )
    h(doc, "5.4 Apply workflow used", 2)
    numbered(
        doc,
        [
            "terraform init in the environment folder.",
            "Local terraform.tfvars (gitignored) with alert_email and groq_api_key.",
            "terraform plan then terraform apply.",
            "Outputs captured: alb_url, ecr_* urls, ecs_cluster_name, grafana_url (staging).",
        ],
    )

    # 6
    h(doc, "6. Containerization & First Deploy", 1)
    bullets(
        doc,
        [
            "docker/server/Dockerfile — multi-stage Node production image.",
            "docker/client/Dockerfile — Vite production build served by nginx.",
            "Images built for linux/amd64 (Fargate).",
            "Logged into ECR, pushed :latest tags, forced new ECS deployments.",
            "Verified /api/health returned OK and frontend returned HTTP 200.",
        ],
    )
    p(doc, "Dev ALB URL:", bold=True)
    p(doc, "http://ai-studyplan-dev-alb-498635700.us-east-1.elb.amazonaws.com")
    p(doc, "Staging ALB URL:", bold=True)
    p(doc, "http://ai-studyplan-staging-alb-524443134.us-east-1.elb.amazonaws.com")

    h(doc, "6.1 Critical runtime fix — API key environment variable", 2)
    p(
        doc,
        "After first deploy, the UI reported: OPENAI_API_KEY environment variable is not set. "
        "Root cause: Terraform initially injected GROQ_API_KEY, but server/services/openaiService.js "
        "reads OPENAI_API_KEY. Fix: map Secrets Manager secret to OPENAI_API_KEY and set "
        "OPENAI_BASE_URL and OPENAI_MODEL for Groq. ECS task definition was updated and redeployed."
    )

    # 7
    h(doc, "7. CI/CD Pipeline", 1)
    p(doc, "GitHub Actions workflows live under .github/workflows/. Guide: .github/CI-CD.md")
    table(
        doc,
        ["Workflow", "Trigger", "What it does"],
        [
            ["deploy-dev.yml", "Push to main (client/server/docker) or manual", "Build images, push ECR, update ECS, wait stable, health check"],
            ["deploy-staging.yml", "Push to staging branch or manual", "Same pattern targeting staging resources"],
            ["deploy-prod.yml", "Manual only; confirm phrase deploy-prod", "Production deploy of a staging-tested git SHA"],
            ["terraform-plan.yml", "PRs touching infra/", "terraform plan and comment on PR"],
            ["reusable-ecs-deploy.yml", "Called by other workflows", "Shared build/push/deploy logic"],
        ],
    )
    h(doc, "7.1 Promotion model", 2)
    numbered(
        doc,
        [
            "Developer merges to main → DEV deploys automatically.",
            "Team tests on DEV, notes commit SHA.",
            "Promote to STAGING (manual workflow or staging branch) and validate.",
            "Only then run Deploy Production with the same SHA and type deploy-prod.",
            "Recommended: GitHub Environment “production” with required reviewers.",
        ],
    )
    h(doc, "7.2 CI credentials issue and resolution", 2)
    p(
        doc,
        "First Actions run failed with: Credentials could not be loaded. Cause: missing repository secrets. "
        "After AWS_ACCESS_KEY_ID and AWS_SECRET_ACCESS_KEY were added in GitHub Settings → Secrets, "
        "the workflow was re-run successfully (AWS login, ECR push, ECS deploy, health checks all green)."
    )
    p(doc, "Note: Node 20 deprecation messages from Actions are warnings only, not the failure cause.")

    # 8
    h(doc, "8. Staging Environment", 1)
    p(
        doc,
        "A full staging stack was created (separate VPC CIDR 10.30.0.0/16, its own ALB/ECR/ECS/EFS/budgets) "
        "so releases can be tested before production. Images were built and pushed to staging ECR and "
        "services force-deployed."
    )

    # 9
    h(doc, "9. Observability — Prometheus, Grafana, Alerts", 1)
    p(doc, "Implemented on staging as an ECS multi-container monitoring task.")
    table(
        doc,
        ["Component", "Role"],
        [
            ["YACE (Yet Another CloudWatch Exporter)", "Exports ECS/ALB CloudWatch metrics to Prometheus"],
            ["Prometheus", "Scrapes YACE; evaluates alert rules"],
            ["Grafana", "Dark executive dashboard: Reliability Command Center"],
            ["Alertmanager", "Included in stack (non-critical if stopped); primary email via SNS"],
            ["CloudWatch metric alarms", "CPU high, ALB 5xx, unhealthy hosts → SNS email"],
        ],
    )
    p(doc, "Grafana URL:", bold=True)
    p(doc, "http://ai-studyplan-staging-alb-524443134.us-east-1.elb.amazonaws.com/grafana/")
    p(doc, "Default admin user: admin (password set via Terraform variable grafana_admin_password — change after first login).")
    p(doc, "Dashboard highlights: healthy/unhealthy targets, ECS CPU/memory, request rate, 5xx, latency, scaling playbook banner.")

    # 10
    h(doc, "10. Cost Management", 1)
    bullets(
        doc,
        [
            "AWS Budgets with email notifications at 80% and 100% of monthly limit.",
            "SNS topic subscription to the project alert email.",
            "CloudWatch EstimatedCharges alarm in us-east-1.",
            "Cost Explorer is a Billing console analysis tool — not provisioned as Terraform resources; "
            "used for reporting by service.",
            "Account MTD spend may include unrelated services (e.g. EKS/WAF). Filter by Project/Environment tags "
            "for this app. Billing data can lag 24–48 hours.",
        ],
    )
    p(doc, "Rough estimate for always-on dev + staging: approximately USD 75–115 per month (ALB + Fargate dominate).")

    # 11
    h(doc, "11. Scaling Behavior", 1)
    p(
        doc,
        "Scaling is ECS service autoscaling (task count), not EC2 Auto Scaling Groups. Server service tracks "
        "average CPU (~60% target) between configured min and max. To scale further, update Terraform variables "
        "server_min_count / server_max_count / desired counts, or temporarily aws ecs update-service --desired-count."
    )
    p(
        doc,
        "For sudden growth (example: hundreds → hundreds of thousands of users), more tasks alone are insufficient. "
        "Plan RDS instead of SQLite, caching, CDN, async AI queues, and stronger ALB request-based autoscaling."
    )

    # 12
    h(doc, "12. Disaster Recovery — Current State", 1)
    table(
        doc,
        ["In place", "Not yet"],
        [
            ["Multi-AZ subnets / ALB / EFS mounts", "Multi-region active failover"],
            ["Terraform rebuild capability", "Formal RTO/RPO runbooks & drills"],
            ["Secrets Manager for API keys", "Automated DB backup strategy for large scale"],
        ],
    )
    p(doc, "DR is basic/multi-AZ oriented, not full enterprise multi-region DR.")

    # 13
    h(doc, "13. Frontend Redesign", 1)
    p(
        doc,
        "The original UI leaned on purple/neon/glass card patterns. It was redesigned into a brand-first, "
        "light-first, three-dimensional study workspace."
    )
    bullets(
        doc,
        [
            "Typography: Fraunces (display) + Sora (body).",
            "Palette: ink, bone, teal (#1faa8a), copper accents — avoiding purple neon defaults.",
            "Hero: brand name AI Study Hub as primary signal; CSS 3D floating planes and orbit rings.",
            "Motion: cursor parallax, dual orbit, pulse lines, CTA sheen, scroll-reveal tool panels, tilt cards, chat message entrance.",
            "Navbar and chat bubbles restyled to match; Notes step indicator colors updated away from purple gradients.",
            "Default theme set to light.",
        ],
    )
    p(doc, "User confirmed: “ok we'll go with this one” — this UI is the adopted product look.")

    # 14
    h(doc, "14. Documentation Added to the Repository", 1)
    bullets(
        doc,
        [
            "PROJECT_DOCUMENTATION.md — narrative overview of the full journey.",
            "infra/README.md — Terraform usage.",
            ".github/CI-CD.md — pipeline and secrets.",
            "infra/MONITORING.md — monitoring image build notes.",
            "docs/AI-StudyHub-Complete-Project-Report.docx — this detailed Word report.",
        ],
    )

    # 15
    h(doc, "15. Key Git Commits / Deliverables (representative)", 1)
    bullets(
        doc,
        [
            "ECS Terraform + Docker packaging for AWS deploy.",
            "GitHub Actions CI/CD for ECS deploy and Terraform plan.",
            "Staging environment + gated prod promotion workflows.",
            "OPENAI_API_KEY / Groq base URL wiring fix on ECS.",
            "Prometheus/Grafana/YACE monitoring on staging + SNS scaling alarms.",
            "Frontend 3D study-plane redesign + animation pass.",
            "Project documentation markdown and this Word report.",
        ],
    )

    # 16
    h(doc, "16. How to Operate Going Forward", 1)
    h(doc, "16.1 Deploy application (DEV)", 2)
    p(doc, "Push to main affecting client/, server/, or docker/ — or run Actions → Deploy Dev (ECS).")
    h(doc, "16.2 Promote to STAGING", 2)
    p(doc, "Actions → Deploy Staging → Run workflow (optionally pass the DEV-tested git SHA).")
    h(doc, "16.3 Production (when ready)", 2)
    numbered(
        doc,
        [
            "Apply infra/environments/prod with Terraform.",
            "Set PROD_ALB_URL / GitHub Environment production with reviewers.",
            "Run Deploy Production with git_ref + confirm=deploy-prod.",
        ],
    )
    h(doc, "16.4 Secrets hygiene", 2)
    bullets(
        doc,
        [
            "Never commit terraform.tfvars containing API keys.",
            "Rotate Groq key if it was shared in chat.",
            "Change Grafana admin password after first login.",
            "Confirm SNS email subscription(s) for budgets and alarms.",
        ],
    )

    # 17
    h(doc, "17. Recommended Next Steps", 1)
    numbered(
        doc,
        [
            "Review one full month in Cost Explorer filtered by project tags; adjust budgets.",
            "Apply production environment when business-ready; keep manual promotion only.",
            "Add ALB RequestCountPerTarget autoscaling policies for spikier traffic.",
            "Migrate SQLite → RDS PostgreSQL before large concurrent user growth.",
            "Extend frontend visual language consistently across Notes and Quiz detail flows.",
            "Formalize backup/restore and a simple DR checklist.",
        ],
    )

    # 18
    h(doc, "18. Conclusion", 1)
    p(
        doc,
        "From an existing AI study application, the team delivered a complete AWS delivery platform: "
        "Terraform-managed ECS infrastructure, multi-stage environments, working CI/CD, cost alerting, "
        "staging observability with Grafana, and a distinctive animated frontend. Development and staging "
        "are live; production remains intentionally gated so releases must pass staging and human confirmation."
    )

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_hr(footer)
    end = doc.add_paragraph()
    end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = end.add_run("— End of Report —\nAI Study Hub · Waleedcoderarch/AI-StudyPlan · us-east-1")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x5B, 0x6B, 0x7C)

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
